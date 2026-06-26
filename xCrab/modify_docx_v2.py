from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

# Load decrypted document
doc = Document('搬迁事宜告知函_decrypted.docx')

# Create a new document to store modified content
new_doc = Document()

# Copy styles from original document
for style in doc.styles:
    try:
        new_doc.styles.add_style(style.name, style.type)
    except:
        pass

# Process paragraphs
paragraphs_to_process = []
early_move_reward_added = False

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text
    
    # 1. Modify second instance content
    if "但我司绝不接受本次不公判决，目前已准备上诉、坚持二审维权，将与房东斗争到底。" in text:
        text = text.replace(
            "但我司绝不接受本次不公判决，目前已准备上诉、坚持二审维权，将与房东斗争到底。",
            "但我司绝不接受本次不公判决，目前已准备上诉。不管二审结果如何，均以租户无关。"
        )
    
    if "租户全部平稳搬离后，我司将留守场地对家具家电和装修覆着物进行清场、坚持二审上诉、继续与房东维权斗争到底，誓死捍卫我方合法权益。" in text:
        text = text.replace(
            "租户全部平稳搬离后，我司将留守场地对家具家电和装修覆着物进行清场、坚持二审上诉、继续与房东维权斗争到底，誓死捍卫我方合法权益。",
            "租户全部平稳搬离后，我司将留守场地对家具家电和装修覆着物进行清场。"
        )
    
    # 2. Add early move reward after "押金全额无条件退还" section
    if "押金全额无条件退还" in text and not early_move_reward_added:
        # Add the current paragraph
        paragraphs_to_process.append(text)
        
        # Add early move reward section
        early_move_reward_text = [
            "早搬奖励",
            "为鼓励各位租户提前搬迁，凡在2026年6月15日前完成搬离并配合交房的租户，可额外享受：",
            "① 每户给予 ¥XXX元搬家补贴（可用于搬家公司或报销搬家费用）；",
            "② 公司可为有需要的租户提供基本的周边房源信息及协助对接；",
            "③ 凡在6月15日前完成搬离并配合交房的租户，可额外免除部分水费/额外补贴。"
        ]
        
        for reward_text in early_move_reward_text:
            paragraphs_to_process.append(reward_text)
        
        early_move_reward_added = True
    else:
        paragraphs_to_process.append(text)

# Add paragraphs to new document
for text in paragraphs_to_process:
    if text.strip():  # Only add non-empty paragraphs
        new_para = new_doc.add_paragraph()
        new_para.text = text

# Save modified document
new_doc.save('搬迁事宜告知函_修改版.docx')
print("Document saved as: 搬迁事宜告知函_修改版.docx")

# Also save with the original filename for easy download
new_doc.save('搬迁事宜告知函_修订标记版_v3.docx')
print("Document also saved as: 搬迁事宜告知函_修订标记版_v3.docx")